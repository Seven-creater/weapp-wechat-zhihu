from __future__ import annotations

from pathlib import Path
import shutil
import textwrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


REPO_ROOT = Path(__file__).resolve().parents[1]
DOC_DIR = Path(r"C:\Users\29785\Desktop\unimportant")
PICTURE_DIR = Path(r"C:\Users\29785\Desktop\pictures")
TARGET_DOC = next(p for p in DOC_DIR.glob("*项目申报书4.20.docx") if not p.name.startswith("~$"))
BACKUP_DOC = DOC_DIR / "【智解行忧_开创老旧小区无障碍环境智诊新纪元】项目申报书4.20.backup-before-2.2-2.3.docx"


def find_picture(prefix: str) -> Path:
    matches = sorted(PICTURE_DIR.glob(f"{prefix}*"))
    if not matches:
        raise FileNotFoundError(f"Missing picture with prefix {prefix!r}")
    return matches[0]


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path(r"C:\Windows\Fonts\msyhbd.ttc") if bold else Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf") if bold else Path(r"C:\Windows\Fonts\simsun.ttc"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def wrap_text(text: str, width: int) -> str:
    return "\n".join(textwrap.wrap(text, width=width, break_long_words=True, replace_whitespace=False))


def draw_box(draw: ImageDraw.ImageDraw, box, text: str, *, fill: str, outline: str, font, text_fill="#10243e"):
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=3)
    wrapped = wrap_text(text, 14)
    bbox = draw.multiline_textbbox((0, 0), wrapped, font=font, spacing=8, align="center")
    tw = bbox[2] - bbox[0]
    th = bbox[3] - bbox[1]
    tx = x1 + (x2 - x1 - tw) / 2
    ty = y1 + (y2 - y1 - th) / 2
    draw.multiline_text((tx, ty), wrapped, font=font, fill=text_fill, spacing=8, align="center")


def draw_arrow(draw: ImageDraw.ImageDraw, start, end, color="#3563e9", width=6):
    draw.line([start, end], fill=color, width=width)
    ex, ey = end
    sx, sy = start
    dx = ex - sx
    dy = ey - sy
    if abs(dx) >= abs(dy):
        sign = 1 if dx >= 0 else -1
        tip = (ex, ey)
        wing1 = (ex - 18 * sign, ey - 10)
        wing2 = (ex - 18 * sign, ey + 10)
    else:
        sign = 1 if dy >= 0 else -1
        tip = (ex, ey)
        wing1 = (ex - 10, ey - 18 * sign)
        wing2 = (ex + 10, ey - 18 * sign)
    draw.polygon([tip, wing1, wing2], fill=color)


def make_canvas(title: str, subtitle: str):
    img = Image.new("RGB", (1800, 1080), "#f7f9fd")
    draw = ImageDraw.Draw(img)
    title_font = load_font(40, bold=True)
    sub_font = load_font(22)
    draw.rounded_rectangle((64, 56, 280, 104), radius=24, fill="#eaf1ff", outline="#b7ccff", width=2)
    draw.text((86, 68), "基于真实代码整理", font=load_font(22, bold=True), fill="#295fe9")
    draw.text((64, 136), title, font=title_font, fill="#122033")
    draw.text((64, 196), subtitle, font=sub_font, fill="#607085")
    return img, draw


def generate_main_function_framework(out_path: Path):
    img, draw = make_canvas(
        "主要功能框架图",
        "围绕微信小程序当前真实能力整理，包括问题上报、地图展示、社区互动、专业方案、施工跟进和消息联动。"
    )
    title_font = load_font(26, bold=True)
    box_font = load_font(24, bold=True)
    small_font = load_font(20)

    draw.rounded_rectangle((520, 320, 1280, 520), radius=24, fill="#ffffff", outline="#ccd7ea", width=3)
    draw.text((770, 350), "微信小程序平台", font=title_font, fill="#13315c")
    draw.text((630, 402), "前端页面 + 云函数 + 云数据库/云存储", font=load_font(28, bold=True), fill="#2c63ea")

    boxes = [
        ((120, 320, 440, 460), "问题上报\n图片上传\n位置绑定", "#e9f7ef", "#7cc89a"),
        ((120, 560, 440, 700), "社区地图\n点位渲染\n附近检索", "#eef4ff", "#90b2ff"),
        ((540, 620, 900, 780), "社区圈列表\n帖子详情\n评论点赞收藏", "#fff4e8", "#f2bf79"),
        ((1360, 320, 1680, 460), "设计方案提交\n专业建议补充", "#f5ecff", "#bb8ef2"),
        ((1360, 560, 1680, 700), "施工项目创建\n节点更新\n进度跟踪", "#ffecef", "#ef9aa9"),
        ((940, 620, 1280, 780), "消息聚合\n局部实时监听\n身份切换", "#eaf6ff", "#74b9ff"),
    ]
    for box, text, fill, outline in boxes:
        draw_box(draw, box, text, fill=fill, outline=outline, font=box_font)

    center_left = (520, 420)
    center_right = (1280, 420)
    center_bottom = (900, 520)
    draw_arrow(draw, (440, 390), center_left)
    draw_arrow(draw, (440, 630), (520, 460))
    draw_arrow(draw, center_right, (1360, 390))
    draw_arrow(draw, (1280, 460), (1360, 630))
    draw_arrow(draw, center_bottom, (720, 620))
    draw_arrow(draw, center_bottom, (1110, 620))

    draw.rounded_rectangle((460, 860, 1340, 980), radius=20, fill="#ffffff", outline="#ccd7ea", width=3)
    draw.text((520, 890), "统一支撑能力：角色权限校验、字段校验、白名单查询、批量取数、GeoPoint 地图检索", font=small_font, fill="#44556b")
    img.save(out_path)


def generate_db_er(out_path: Path):
    img, draw = make_canvas(
        "云数据库 ER 图",
        "依据当前小程序中真实使用的 CloudBase 集合关系整理，突出 posts 与用户、互动、专业服务之间的关联。"
    )
    box_font = load_font(20)

    boxes = {
        "users": (120, 330, 450, 520),
        "posts": (730, 270, 1070, 520),
        "comments": (120, 620, 450, 780),
        "actions": (480, 620, 810, 780),
        "design_proposals": (1090, 620, 1500, 780),
        "construction_projects": (1140, 300, 1680, 500),
        "messages": (1480, 620, 1760, 780),
    }
    texts = {
        "users": "users\n_openid\nuserType\nuserInfo\nisAdmin",
        "posts": "posts\n_id\n_openid\ntype=status\ncontent/images\ncommunity/location\naiSolution/stats",
        "comments": "comments\npostId\nauthorOpenid\ncontent\ncreateTime",
        "actions": "actions\npostId/targetId\ntype\n_openid\ncreateTime",
        "design_proposals": "design_proposals\npostId / issueId\ndesignerId\ncontent/status",
        "construction_projects": "construction_projects\nissueId\ncontractorId\nstatus/stages\nprogress nodes",
        "messages": "messages\nchatId\nsenderId\ncontent\ncreateTime",
    }
    fills = {
        "users": "#eef4ff",
        "posts": "#ffffff",
        "comments": "#fff5e8",
        "actions": "#fff5e8",
        "design_proposals": "#f3ecff",
        "construction_projects": "#ffecef",
        "messages": "#eaf7ef",
    }
    outlines = {
        "users": "#86a9ff",
        "posts": "#b8c5d8",
        "comments": "#f1c27d",
        "actions": "#f1c27d",
        "design_proposals": "#c59af5",
        "construction_projects": "#ee9cad",
        "messages": "#7dc999",
    }
    for key, box in boxes.items():
        draw_box(draw, box, texts[key], fill=fills[key], outline=outlines[key], font=box_font)

    def c(box):
        x1, y1, x2, y2 = box
        return ((x1 + x2) // 2, (y1 + y2) // 2)

    draw_arrow(draw, (450, 410), (730, 370))
    draw_arrow(draw, (450, 690), (730, 450))
    draw_arrow(draw, (810, 690), (900, 520))
    draw_arrow(draw, (1070, 430), (1140, 390))
    draw_arrow(draw, (1070, 480), (1240, 620))
    draw_arrow(draw, (450, 410), (730, 470))
    draw_arrow(draw, (450, 410), (480, 650))
    draw_arrow(draw, (450, 410), (1500, 690))

    label_font = load_font(18, bold=True)
    draw.text((540, 334), "_openid", font=label_font, fill="#3563e9")
    draw.text((535, 520), "postId", font=label_font, fill="#c06c00")
    draw.text((900, 560), "issueId", font=label_font, fill="#b03b66")
    draw.text((1080, 560), "postId/issueId", font=label_font, fill="#7b4dca")
    draw.text((560, 448), "_openid", font=label_font, fill="#3563e9")
    img.save(out_path)


def generate_login_sequence(out_path: Path):
    img, draw = make_canvas(
        "登录与身份初始化时序图",
        "根据 login 页面、全局登录流程和 users 集合初始化逻辑整理，展示首次进入小程序时的真实调用链路。"
    )
    header_font = load_font(24, bold=True)
    body_font = load_font(20)
    x_positions = [180, 560, 940, 1320, 1640]
    headers = ["用户", "微信小程序", "login 云函数", "users 集合", "全局状态"]
    for x, name in zip(x_positions, headers):
        draw.rounded_rectangle((x - 90, 290, x + 90, 350), radius=16, fill="#ffffff", outline="#cdd8ea", width=3)
        bbox = draw.textbbox((0, 0), name, font=header_font)
        draw.text((x - (bbox[2] - bbox[0]) / 2, 304), name, font=header_font, fill="#123058")
        draw.line((x, 350, x, 900), fill="#b7c6df", width=3)

    steps = [
        (390, 180, 560, "打开小程序"),
        (450, 560, 940, "调用 login 获取 openid"),
        (520, 940, 1320, "查询用户资料"),
        (590, 1320, 940, "返回用户记录或空结果"),
        (660, 940, 1320, "首次用户则创建资料"),
        (730, 940, 560, "返回 openid / userType"),
        (800, 560, 1640, "写入全局登录态与用户信息"),
        (870, 1640, 560, "页面按角色渲染入口"),
    ]
    for y, sx, ex, label in steps:
        draw_arrow(draw, (sx, y), (ex, y))
        bbox = draw.textbbox((0, 0), label, font=body_font)
        draw.rounded_rectangle((min(sx, ex) + 20, y - 28, min(sx, ex) + 20 + (bbox[2] - bbox[0]) + 18, y - 2), radius=10, fill="#ffffff", outline="#d9e2ec")
        draw.text((min(sx, ex) + 30, y - 24), label, font=body_font, fill="#44556b")

    draw.rounded_rectangle((150, 940, 1650, 1010), radius=18, fill="#ffffff", outline="#ccd7ea", width=2)
    draw.text((200, 962), "结果：用户完成微信登录后，平台在 users 集合中识别或初始化身份信息，后续页面再结合角色权限控制功能入口。", font=body_font, fill="#44556b")
    img.save(out_path)


def generate_marker_code_figure(out_path: Path):
    snippet_path = REPO_ROOT / "pages" / "index" / "index.js"
    lines = snippet_path.read_text(encoding="utf-8").splitlines()
    left = "\n".join(lines[319:357])
    right = "\n".join(lines[357:390])

    img, draw = make_canvas(
        "多个帖子信息映射到地图点位的核心代码图",
        "依据 pages/index/index.js 中的真实实现整理，展示 issue 帖子如何提取位置并组装为 map markers。"
    )

    mono_font = load_font(17)
    title_font = load_font(24, bold=True)

    panels = [
        ((80, 300, 860, 980), "位置提取：extractPostLocation()", left),
        ((940, 300, 1720, 980), "点位组装：buildDisplayMarkers()", right),
    ]
    for box, title, code in panels:
        x1, y1, x2, y2 = box
        draw.rounded_rectangle(box, radius=18, fill="#ffffff", outline="#cfd9ea", width=3)
        draw.rounded_rectangle((x1 + 18, y1 + 18, x2 - 18, y1 + 68), radius=12, fill="#eef4ff", outline="#c4d6ff", width=2)
        draw.text((x1 + 34, y1 + 30), title, font=title_font, fill="#1f4fd8")
        code_box = (x1 + 18, y1 + 90, x2 - 18, y2 - 18)
        draw.rounded_rectangle(code_box, radius=14, fill="#111827", outline="#293244", width=2)
        margin_x = code_box[0] + 20
        margin_y = code_box[1] + 18
        for idx, line in enumerate(code.splitlines()):
            y = margin_y + idx * 18
            if y > code_box[3] - 24:
                break
            draw.text((margin_x, y), line[:88], font=mono_font, fill="#eaf0f8")

    draw.rounded_rectangle((310, 190, 1490, 250), radius=16, fill="#ffffff", outline="#d7e0ee", width=2)
    draw.text((356, 208), "真实路径：getPublicData -> normalizeIssuePost -> extractPostLocation -> buildDisplayMarkers -> map(markers)", font=load_font(20, bold=True), fill="#4a5a70")
    img.save(out_path)


def set_paragraph_text(paragraph: Paragraph, text: str):
    style = paragraph.style
    alignment = paragraph.alignment
    while paragraph.runs:
        paragraph._p.remove(paragraph.runs[0]._r)
    if text:
        run = paragraph.add_run(text)
        run.font.size = Pt(11)
    paragraph.style = style
    paragraph.alignment = alignment


def add_caption_before(anchor: Paragraph, text: str):
    p = anchor.insert_paragraph_before(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p.runs:
        p.runs[0].font.size = Pt(10.5)
    return p


def add_image_before(anchor: Paragraph, image_path: Path, width_inches: float = 5.9):
    p = anchor.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    return p


def find_paragraph_index_by_text(doc: Document, text: str) -> int:
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == text:
            return idx
    raise ValueError(f"Paragraph not found: {text}")


def replace_table_image(table, image_path: Path, row: int = 0, col: int = 0, width_inches: float = 6.1):
    cell = table.cell(row, col)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))


def strip_text_from_image_cell(cell):
    paragraphs = list(cell.paragraphs)
    for paragraph in paragraphs:
        has_drawing = bool(paragraph._p.xpath(".//w:drawing"))
        if has_drawing:
            for node in paragraph._p.xpath(".//w:t"):
                node.text = ""
        elif paragraph.text.strip():
            paragraph._element.getparent().remove(paragraph._element)
    if not cell.paragraphs:
        cell.text = ""


def prepare_extra_images():
    PICTURE_DIR.mkdir(parents=True, exist_ok=True)
    extra = {
        "10-main-function-framework.png": generate_main_function_framework,
        "11-cloudbase-db-er.png": generate_db_er,
        "12-login-sequence.png": generate_login_sequence,
        "13-map-marker-code.png": generate_marker_code_figure,
    }
    for name, builder in extra.items():
        path = PICTURE_DIR / name
        builder(path)


def main():
    prepare_extra_images()

    if BACKUP_DOC.exists():
        shutil.copyfile(BACKUP_DOC, TARGET_DOC)

    doc = Document(TARGET_DOC)

    replacements = {
        112: "第二步，推进空间定位与数据化呈现，搭建单体诊断与空间地图的融合基础。在完成问题识别与上报流程设计后，团队进一步引入空间地图能力，依托微信小程序与云开发实现障碍点信息的采集、存储、检索与可视化展示，具体开发过程包括以下四个环节：",
        114: "① 小程序端采集与空间绑定——居民在问题上报页面拍摄或选择现场图片，填写问题描述、问题类别和所属社区信息；页面同时支持通过 wx.chooseLocation 手动选点，并在需要时回退使用 wx.getLocation 获取当前位置，经纬度、地址、图片、文本描述和用户建议等字段在前端统一组装为上传对象。页面还预留了 AI 辅助分析入口，可生成 aiSolution 作为后续查看与协同处理的参考内容。",
        116: "② 数据上传与云存储——小程序先通过 wx.cloud.uploadFile 将现场图片上传至云存储，获得对应的 fileID 列表；随后再通过 wx.cloud.callFunction 调用 createIssuePost 云函数，将位置坐标、图片标识、问题描述、社区名称、类别信息和辅助分析结果统一写入云端。云函数内部会将前端传入的经纬度转换为云数据库支持的 GeoPoint 类型，并完成字段合法性校验。",
        117: "③ 数据库设计——当前项目实际采用微信云开发 CloudBase 文档型数据库存储障碍点信息，而非 PostgreSQL/PostGIS。以 posts 集合中的问题上报数据为例，单条记录主要包含 _id、_openid、type、status、title、content、images、category/categoryId/categoryName、community、location（GeoPoint）、address/formattedAddress/detailAddress、userSuggestion、aiSolution、stats、createTime 和 updateTime 等字段，能够覆盖问题上报、地图展示与后续协同处理所需的核心信息。",
        118: "④ 空间检索与可视化——地图首页通过云函数 getPublicData 传入中心点与检索半径，在后端使用 geoNear 实现附近问题帖的快速查询；前端获得数据后进一步提取 GeoPoint 中的经纬度信息，构建 markers 数组并交由微信地图组件渲染，实现社区障碍点在地图上的分布展示。用户点击点位后可继续进入帖子详情查看问题描述、图片及互动信息，为后续社区级问题研判与改造统筹提供直观的数据支撑。",
        120: "相关小程序采集流程图、云数据库字段结构图和社区点位分布渲染图如下所示：",
        122: "第三步，推进社区级问题汇聚与协同研判，形成从点位上报到后续处置的联动基础。在完成空间定位与数据化呈现后，项目进一步依托社区地图、问题详情、设计方案和施工项目等模块，实现对障碍点问题的集中查看、关联处理与阶段跟踪。",
        123: "多个帖子信息映射到地图点位的核心代码如下图所示：",
        124: "在前端实现上，地图页会根据当前位置或指定社区范围拉取附近的 issue 数据，并将点位与帖子内容、图片、地址信息和互动状态统一展示；社区工作者、设计师和施工方可以据此快速查看问题集中区域，进入详情页进一步了解现场情况和已有回应内容，从而形成以空间分布为线索的问题研判方式。",
        126: "在业务联动上，帖子详情页还可以继续衔接设计方案提交、施工项目创建和进度更新等操作，设计方案数据写入 design_proposals 集合，施工过程数据写入 construction_projects 集合，问题点位、专业建议与施工进展由此形成可追踪的关联链路。虽然当前版本尚未实现基于空间数据库的自动聚类排序和标准化报告一键生成，但已经完成了问题采集、空间展示、专业响应和进度跟踪的数字化闭环，为后续社区级诊断分析与报告生成提供了可靠的数据基础。",
        127: "",
        128: "通过上述三步逐步推进，项目已完成从问题识别、空间定位到社区级协同处理的基础链路建设，实现了单点问题记录、地图展示与后续专业服务衔接的全流程数字化，为平台继续扩展社区级诊断分析与改造统筹能力奠定了技术基础。",
        131: "项目开发的第三步核心工作，是搭建供需协同数字平台，打通居民、社区工作者、设计师和施工方等多方信息联动链路，实现问题上报、需求流转、专业建议和施工跟进的线上协同。平台整体采用微信小程序原生框架与 CloudBase 云开发一体化架构，前端负责页面交互与地图展示，云函数负责业务校验和数据读写，云数据库与云存储分别承载结构化业务数据和图片文件。基于当前项目代码实现，平台建设过程可概括为以下十一个步骤：",
        132: "第一步，搭建项目基础架构，明确小程序端、云函数和云开发资源之间的分工。前端页面负责地图、社区、发布、消息和个人中心等交互入口，核心业务逻辑统一放在云函数中处理，数据分别存储于云数据库和云存储中；项目目录按 pages、cloudfunctions、utils 等模块划分，便于后续维护与协作开发。",
        134: "第二步，配置页面路由与基础导航。在 app.json 中统一注册页面路径，并设置“地图、社区、发布、消息、我的”五个底部导航入口，使用户能够快速进入核心功能页面；不同角色的专属能力则在统一导航基础上再结合权限控制逐步开放。",
        140: "第三步，完成登录与统一云调用封装。用户首次进入小程序后，通过登录流程获取 openid，并在 users 集合中查询或初始化个人资料和角色信息；同时在全局层封装统一的云函数调用方法，供各页面复用，从而减少重复请求处理逻辑并保证错误提示的一致性。",
        143: "第四步，设计核心数据集合，明确平台的业务数据承载方式。结合当前实现，平台主要围绕以下集合组织用户、帖子、互动记录和专业服务数据：",
        150: "各集合之间主要通过 _openid、postId、issueId 等字段建立关联，以支撑帖子详情、互动消息、设计方案和施工项目之间的联动查询。",
        153: "第五步，建立角色权限控制机制，实现不同身份的差异化功能适配。在 utils/userTypes.js 中定义普通居民、设计师、施工方、社区工作者等角色，并为每类角色配置对应的能力边界，例如设计师可提交方案，施工方可创建项目，社区工作者可查看和处理本社区相关信息。",
        155: "前端页面根据当前用户角色采用条件渲染控制按钮和入口显示，后端云函数在执行写操作前再次校验调用者身份，从而形成“前端可见范围控制 + 后端权限兜底”的双层保障机制，降低越权调用风险。",
        157: "第六步，实现居民问题上报功能。居民在发布页面填写问题描述、选择类别和社区、上传现场图片并选择位置后，前端先将图片上传至云存储，再把文本、类别、位置和图片 fileID 一并提交给 createIssuePost 云函数。云函数完成必填项、长度和位置字段校验后，将问题写入 posts 集合并标记为 issue 类型，形成后续社区协同的基础数据。",
        159: "第七步，开发社区圈列表与详情互动，形成问题展示和持续跟进入口。",
        160: "社区圈列表页调用 getPublicData 云函数，传入分页参数、排序方式和字段白名单，只返回卡片展示所需的最小字段集合，以控制响应体积并提升加载效率；云函数在后端限制可查询集合范围，避免前端随意透传集合名。",
        161: "帖子详情页展示完整内容、AI 辅助分析结果、评论区和点赞收藏入口。用户提交评论时，前端调用 createComment 云函数，后端会先校验帖子是否存在，再写入 comments 集合并同步更新 posts 中的评论统计，从而实现基础互动闭环。",
        163: "第八步，接入设计师与施工方服务，推动居民问题向专业处置环节流转。设计师可在帖子详情页提交改造方案，施工方可基于问题创建施工项目；相关数据分别写入 design_proposals 与 construction_projects 集合，并与 issue 帖子建立关联，便于后续查看方案内容、项目状态和节点进展。",
        164: "第九步，实现消息聚合与局部实时同步。用户打开“消息”页面时，前端调用 getNotificationFeed 云函数，从 comments、actions 和 posts 等集合中聚合与本人帖子相关的评论与点赞动态，并按时间倒序展示。",
        165: "在实时同步方面，当前版本已在聊天页面和部分详情页面接入云数据库 watch 监听，用于感知消息流或个别业务数据变化；当有新的聊天消息或相关数据更新时，前端页面可自动刷新，从而提升多方协同时效性。",
        166: "第十步，加强安全校验与性能优化。当前工程实现中，首先通过 _shared/validate.js 对字符串、枚举、ID 列表等输入进行统一校验，控制字段类型、长度与必填规则，减少脏数据写入风险。",
        167: "其次，在管理员能力控制方面，_shared/auth.js 支持从环境变量和用户资料中识别管理员身份，用于认证审核、内容治理等敏感操作的权限判定，从而避免将高权限能力暴露给普通用户。",
        168: "再次，在性能优化方面，列表页采用请求令牌避免旧请求覆盖新结果，用户资料通过 getUsersBatch 批量补齐，列表与通知查询尽量使用字段投影缩小返回体积，从而降低重复请求和前端渲染压力。",
        169: "第十一步，完成部署与联调。整体联调顺序遵循“先云环境初始化，再核心链路验证，最后扩展角色能力”的原则：",
        180: "通过以上步骤，平台已完成从问题上报、地图展示、社区互动到专业服务和施工跟进的主要业务闭环建设，打通了居民、社区与专业方之间的信息流转链路，具备基层轻量化协同治理所需的核心功能。",
        182: "本项目围绕老旧小区无障碍问题采集、信息协同和专业处置需求，构建了基于微信小程序与云开发的一体化服务平台。平台以“问题上报 - 地图展示 - 社区互动 - 专业响应 - 进度跟踪”为主线，将现场图片、地理位置、帖子互动、设计方案和施工项目等信息统一沉淀到云端，兼顾居民使用的便捷性与社区协同处理的可操作性，为后续社区级诊断分析和改造统筹提供数字化支撑。",
        185: "问题上报与空间定位功能",
        186: "社区互动与专业服务衔接功能",
        187: "项目的主要功能围绕真实业务链路展开。其一，居民可在小程序中完成问题描述填写、图片上传、社区选择和位置标注，生成带有 GeoPoint 坐标的 issue 帖子，并在地图首页查看附近障碍点分布。其二，社区圈提供列表浏览、帖子详情、评论、点赞和收藏等互动能力，支持围绕同一问题持续补充信息。其三，设计师可基于问题帖提交改造方案，施工方可据此创建施工项目并更新节点进度，使问题记录进一步衔接专业建议与实施环节。其四，消息页能够聚合与本人帖子相关的评论和点赞动态，聊天页面与部分详情页面还支持局部实时刷新，从而提升多方协作效率。整体来看，平台已经实现了从问题采集、信息展示到专业处置跟进的核心业务闭环。",
        189: "2.3.2次要功能",
        190: "次要功能主要承担协同支撑、平台治理和后续扩展的作用，既服务于核心业务链路的稳定运行，也为项目后续能力迭代预留了接口基础，具体包括以下两个方面：",
        191: "1. 协同支撑功能：平台支持身份切换与角色识别，不同角色可获得差异化的功能入口；同时通过消息聚合、点赞收藏、评论提醒和局部实时监听等能力，增强居民、社区工作者、设计师与施工方之间的信息同步效率，减少问题流转过程中的沟通断层。",
        192: "2. 平台治理与扩展功能：平台在云函数侧引入统一字段校验、管理员鉴权和集合白名单控制机制，以保障内容安全和数据边界；在业务层则预留了 AI 辅助分析字段、设计方案扩展字段和施工进度字段，便于后续继续完善问题诊断说明、专业方案表达和项目管理能力。",
        193: "总体而言，本项目已形成“采集 - 展示 - 互动 - 协同 - 治理”的完整功能框架。核心功能聚焦老旧小区无障碍问题的数字化记录与协同处置，次要功能则在权限治理、消息联动和后续扩展上提供支撑，使平台兼具现实可用性、可维护性和后续拓展空间。"
    }

    for idx, text in replacements.items():
        set_paragraph_text(doc.paragraphs[idx], text)

    # Capture anchors before inserting any new paragraphs so later insertions don't drift.
    spatial_anchor = doc.paragraphs[121]
    marker_anchor = doc.paragraphs[124]

    # Insert verified mini-program figures for the spatial step right after the "此处可插入" lead-in.
    add_image_before(spatial_anchor, find_picture("07-"))
    add_caption_before(spatial_anchor, "图：小程序采集流程图")
    add_image_before(spatial_anchor, find_picture("08-"))
    add_caption_before(spatial_anchor, "图：云数据库字段结构图")
    add_image_before(spatial_anchor, find_picture("09-"))
    add_caption_before(spatial_anchor, "图：社区点位分布渲染图")

    # Insert the actual map-marker code figure at the former "陈胜文" placeholder.
    add_image_before(marker_anchor, PICTURE_DIR / "13-map-marker-code.png")
    add_caption_before(marker_anchor, "图：多个帖子信息映射到地图点位的核心代码图")

    # python-docx occasionally drops the first image when multiple paragraphs are inserted before the same anchor.
    # If the 07 figure is missing, reinsert it directly between its caption and the next figure block.
    pic07 = find_picture("07-")
    rels = doc.part.rels
    cap07_idx = find_paragraph_index_by_text(doc, "图：小程序采集流程图")
    next_drawing_size = None
    next_drawing_anchor = None
    for idx in range(cap07_idx + 1, len(doc.paragraphs)):
        paragraph = doc.paragraphs[idx]
        blips = paragraph._p.xpath(".//a:blip")
        if blips:
            rid = blips[0].get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
            target = rels[rid]._target
            next_drawing_size = len(target.blob)
            next_drawing_anchor = paragraph
            break
        if paragraph.text.strip() == "图：云数据库字段结构图":
            next_drawing_anchor = paragraph
            break
    if next_drawing_anchor is not None and next_drawing_size != pic07.stat().st_size:
        add_image_before(next_drawing_anchor, pic07)

    # Keep surrounding placeholder paragraphs clean.
    set_paragraph_text(doc.paragraphs[119], "")
    set_paragraph_text(doc.paragraphs[121], "")

    # Replace key verified figures in 2.2.3 and 2.3.
    replace_table_image(doc.tables[17], find_picture("01-"))
    replace_table_image(doc.tables[19], PICTURE_DIR / "12-login-sequence.png")
    replace_table_image(doc.tables[20], PICTURE_DIR / "11-cloudbase-db-er.png")
    replace_table_image(doc.tables[22], find_picture("03-"))
    replace_table_image(doc.tables[23], find_picture("04-"))
    replace_table_image(doc.tables[26], find_picture("06-"))
    replace_table_image(doc.tables[31], PICTURE_DIR / "10-main-function-framework.png")

    # Table 25 carries images but also has stray duplicated text in image rows. Strip the text only.
    for row_idx in (2, 4, 6):
        strip_text_from_image_cell(doc.tables[25].cell(row_idx, 0))

    doc.save(TARGET_DOC)
    print(TARGET_DOC)


if __name__ == "__main__":
    main()
