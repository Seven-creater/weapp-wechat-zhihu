from __future__ import annotations

from pathlib import Path
import shutil

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


DOC_DIR = Path(r"C:\Users\29785\Desktop\unimportant")
PIC_DIR = Path(r"C:\Users\29785\Desktop\pictures")
TARGET = DOC_DIR / "2.2.3搭建供需协同平台.docx"
BACKUP = TARGET.with_name(f"{TARGET.stem}.backup-before-condense.docx")


def pic(prefix: str) -> Path:
    matches = sorted(PIC_DIR.glob(f"{prefix}*"))
    if not matches:
        raise FileNotFoundError(f"Missing image with prefix {prefix}")
    return matches[0]


def font(size: int, bold: bool = False):
    candidates = [
        Path(r"C:\Windows\Fonts\msyhbd.ttc") if bold else Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf") if bold else Path(r"C:\Windows\Fonts\simsun.ttc"),
    ]
    for item in candidates:
        if item.exists():
            return ImageFont.truetype(str(item), size=size)
    return ImageFont.load_default()


def arrow(draw: ImageDraw.ImageDraw, start, end, color="#2b63e6"):
    draw.line([start, end], fill=color, width=6)
    ex, ey = end
    sx, sy = start
    sign = 1 if ex >= sx else -1
    draw.polygon([(ex, ey), (ex - 18 * sign, ey - 11), (ex - 18 * sign, ey + 11)], fill=color)


def rounded_text(draw: ImageDraw.ImageDraw, box, title: str, subtitle: str, fill: str, outline: str):
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=3)
    x1, y1, x2, _ = box
    draw.text((x1 + 24, y1 + 26), title, font=font(28, True), fill="#10233f")
    draw.text((x1 + 24, y1 + 72), subtitle, font=font(19), fill="#4f6075")


def generate_compact_issue_flow() -> Path:
    out = PIC_DIR / "14-小程序问题上报简化流程图.png"
    img = Image.new("RGB", (1800, 980), "#f7f9fd")
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle((70, 56, 280, 106), radius=24, fill="#eaf1ff", outline="#bfd1ff", width=2)
    draw.text((94, 68), "真实流程整理", font=font(22, True), fill="#245fe5")
    draw.text((70, 140), "小程序问题上报与云端落库流程", font=font(44, True), fill="#10233f")
    draw.text((70, 208), "基于 pages/issue-edit/index.js 与 cloudfunctions/createIssuePost/index.js 的真实调用链路整理。", font=font(22), fill="#5b6b80")

    boxes = [
        ((80, 350, 380, 520), "填写问题", "描述、类别、社区"),
        ((470, 350, 770, 520), "采集现场信息", "图片 + 位置坐标"),
        ((860, 350, 1160, 520), "上传云存储", "uploadFile 返回 fileID"),
        ((1250, 350, 1550, 520), "调用云函数", "createIssuePost 校验"),
        ((670, 680, 1130, 840), "写入 posts 集合", "GeoPoint、images、content、community、aiSolution"),
    ]
    colors = [("#eef4ff", "#90adff"), ("#eaf8f1", "#7fc99a"), ("#fff4e8", "#f2be76"), ("#f4ecff", "#bf94ee"), ("#ffffff", "#c9d4e6")]
    for (box, title, subtitle), (fill, outline) in zip(boxes, colors):
        rounded_text(draw, box, title, subtitle, fill, outline)

    arrow(draw, (380, 435), (470, 435))
    arrow(draw, (770, 435), (860, 435))
    arrow(draw, (1160, 435), (1250, 435))
    arrow(draw, (1400, 520), (980, 680))

    draw.rounded_rectangle((80, 880, 1720, 930), radius=16, fill="#ffffff", outline="#d6e0ef", width=2)
    draw.text((120, 893), "关键接口：wx.chooseMedia / wx.chooseLocation / wx.getLocation / wx.cloud.uploadFile / wx.cloud.callFunction('createIssuePost')", font=font(20, True), fill="#42546b")
    img.save(out)
    return out


def generate_compact_community_flow() -> Path:
    out = PIC_DIR / "15-社区圈核心数据流简化图.png"
    img = Image.new("RGB", (1800, 980), "#f7f9fd")
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle((70, 56, 280, 106), radius=24, fill="#eaf1ff", outline="#bfd1ff", width=2)
    draw.text((94, 68), "真实流程整理", font=font(22, True), fill="#245fe5")
    draw.text((70, 140), "社区圈核心数据流", font=font(44, True), fill="#10233f")
    draw.text((70, 208), "基于 issue-edit、createIssuePost、getPublicData、post-detail 和 createComment 的真实调用关系整理。", font=font(22), fill="#5b6b80")

    boxes = [
        ((90, 350, 390, 520), "居民上报", "issue-edit 提交问题"),
        ((480, 350, 780, 520), "云端落库", "createIssuePost 写入 posts"),
        ((870, 350, 1170, 520), "列表展示", "getPublicData 分页查询"),
        ((1260, 350, 1560, 520), "详情互动", "post-detail 查看与操作"),
        ((680, 680, 1120, 840), "评论与统计更新", "createComment 写入 comments 并更新 posts.stats"),
    ]
    colors = [("#eef4ff", "#90adff"), ("#eaf8f1", "#7fc99a"), ("#fff4e8", "#f2be76"), ("#f4ecff", "#bf94ee"), ("#ffffff", "#c9d4e6")]
    for (box, title, subtitle), (fill, outline) in zip(boxes, colors):
        rounded_text(draw, box, title, subtitle, fill, outline)

    arrow(draw, (390, 435), (480, 435))
    arrow(draw, (780, 435), (870, 435))
    arrow(draw, (1170, 435), (1260, 435))
    arrow(draw, (1410, 520), (1010, 680))

    draw.rounded_rectangle((90, 880, 1710, 930), radius=16, fill="#ffffff", outline="#d6e0ef", width=2)
    draw.text((130, 893), "核心集合：posts / comments / actions；关键云函数：createIssuePost / getPublicData / createComment", font=font(20, True), fill="#42546b")
    img.save(out)
    return out


def set_normal_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(11)
    for style_name in ["Heading 1", "Heading 2"]:
        style = doc.styles[style_name]
        style.font.name = "黑体"
        style.font.size = Pt(14 if style_name == "Heading 1" else 12)


def add_para(doc: Document, text: str, *, first_line_indent: bool = True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    if first_line_indent:
        p.paragraph_format.first_line_indent = Pt(22)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_image(doc: Document, image_path: Path, caption: str, width: float = 5.9) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    r = cap.add_run(caption)
    r.font.size = Pt(10.5)


def add_section(doc: Document, title: str, body: str, figures: list[tuple[Path, str]]) -> None:
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(6)
    heading.paragraph_format.space_after = Pt(3)
    run = heading.add_run(title)
    run.bold = True
    run.font.size = Pt(12)

    for para in body.split("\n"):
        para = para.strip()
        if para:
            add_para(doc, para)

    for image_path, caption in figures:
        add_image(doc, image_path, caption)


def main() -> None:
    compact_issue_flow = generate_compact_issue_flow()
    compact_community_flow = generate_compact_community_flow()

    if not BACKUP.exists():
        shutil.copyfile(TARGET, BACKUP)

    doc = Document()
    set_normal_style(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("2.2.3 搭建供需协同数字平台")
    r.bold = True
    r.font.size = Pt(16)

    add_para(
        doc,
        "项目第三步聚焦供需协同数字平台建设，目标是将居民问题上报、社区统筹、设计师方案和施工方进度纳入同一条线上协同链路。平台采用微信小程序原生框架与 CloudBase 云开发，前端负责地图、社区、发布、消息和个人中心等交互入口，云函数负责业务校验与数据读写，云数据库和云存储分别承载结构化业务数据和现场图片文件。为避免开发过程描述过于零散，实际建设流程可归纳为以下五个板块。",
    )

    add_section(
        doc,
        "第一板块，确定云开发架构，完成平台基础部署。",
        "团队首先确定以微信小程序作为轻量化入口，并使用 CloudBase 统一承接后端能力。小程序端负责页面展示、表单提交、地图点位渲染和消息入口；云函数负责登录、问题发布、列表查询、评论互动、方案提交、项目创建和通知聚合等业务逻辑；云数据库保存 users、posts、comments、actions、design_proposals、construction_projects 等核心集合；云存储用于保存居民上传的现场图片和方案附件。这样的架构避免了自建服务器成本，也便于后续在微信生态中推广和联调。",
        [(pic("01-"), "图2.5 供需协同数字平台整体技术架构图")],
    )

    add_section(
        doc,
        "第二板块，打通登录、数据模型与角色权限。",
        "平台通过微信登录获取用户 openid，并在 users 集合中维护用户资料、角色类型和认证状态。围绕业务流转，数据库集合之间主要通过 _openid、postId 和 issueId 形成关联：posts 记录居民上报和社区内容，comments 与 actions 保存评论、点赞和收藏，design_proposals 承接设计师方案，construction_projects 承接施工项目和进度节点。\n在权限控制上，前端根据用户角色显示不同按钮和入口，后端云函数在执行设计方案、施工项目、管理操作等写入动作前再次校验身份，形成“前端入口控制 + 后端权限兜底”的双层机制，避免仅依赖页面显示来判断权限。",
        [
            (PIC_DIR / "12-login-sequence.png", "图2.6 登录与身份初始化时序图"),
            (PIC_DIR / "11-cloudbase-db-er.png", "图2.7 云数据库核心集合关系图"),
            (pic("03-"), "图2.8 角色权限校验与前端条件渲染逻辑图"),
        ],
    )

    add_section(
        doc,
        "第三板块，建设问题上报与社区圈互动闭环。",
        "居民端的核心入口是问题上报。用户填写问题描述、选择问题类别和所属社区，拍摄或选择现场图片，并通过位置选择能力绑定经纬度。前端先调用 wx.cloud.uploadFile 将图片上传到云存储，再调用 createIssuePost 云函数写入 posts 集合，帖子类型标记为 issue。\n社区圈列表通过 getPublicData 云函数分页读取帖子数据，并使用字段投影只返回卡片展示所需字段，降低响应体积。进入详情页后，用户可以查看完整内容、图片、AI 辅助分析结果和评论区，也可以进行评论、点赞、收藏等互动；评论写入 comments 集合后，会同步更新帖子统计信息，形成“居民上报 - 社区展示 - 详情互动”的基础闭环。",
        [
            (compact_issue_flow, "图2.9 小程序问题采集与上传绑定流程图"),
            (compact_community_flow, "图2.10 社区圈核心数据流图"),
        ],
    )

    add_section(
        doc,
        "第四板块，衔接设计施工服务，实现专业响应与状态同步。",
        "在居民问题沉淀为 issue 帖子后，平台继续向专业服务侧延伸。设计师可以基于问题详情提交改造建议或设计方案，方案数据写入 design_proposals 集合；施工方可以基于问题创建 construction_projects 项目，并记录项目状态、施工阶段和进度节点。这样，单个问题不再停留在简单反馈层面，而是可以继续关联专业方案和实施进展。\n消息与同步能力用于减少多方沟通断层。消息页通过 getNotificationFeed 聚合与本人帖子相关的评论、点赞等动态；聊天和部分详情页面使用云数据库 watch 监听数据变化，在消息或局部业务状态更新时自动刷新页面，提高居民、社区和专业方之间的协作效率。",
        [
            (pic("05-"), "图2.11 设计方案与施工项目集合关系图"),
            (pic("06-"), "图2.12 消息通知与 watch 监听时序图"),
        ],
    )

    add_section(
        doc,
        "第五板块，完善安全校验、性能优化与联调部署。",
        "平台在云函数侧补充输入校验和权限判断，使用共享校验方法限制字符串长度、枚举值、ID 列表和必填字段，减少脏数据写入；管理员能力通过环境变量和用户表权限共同识别，用于认证审核和内容治理等敏感操作。\n性能方面，社区圈列表使用请求令牌避免旧请求覆盖新结果，用户资料通过批量接口补齐，列表和通知查询尽量使用字段投影缩小返回体积。部署联调按“云环境初始化 - 数据集合创建 - 云函数部署 - 登录验证 - 问题发布 - 社区互动 - 专业服务 - 消息同步”的顺序推进，确保平台主流程稳定可用后，再进行界面细节和后续能力扩展。",
        [(PIC_DIR / "10-main-function-framework.png", "图2.13 供需协同平台主要功能框架图")],
    )

    add_para(
        doc,
        "通过上述五个板块的开发，平台完成了从居民问题采集、社区圈展示、专业服务对接到施工进度跟踪的主要业务链路建设。整体实现既符合微信小程序轻量化使用场景，也为后续社区级诊断分析、改造任务管理和多方协同治理提供了可持续扩展的数字化基础。",
    )

    doc.save(TARGET)
    print(TARGET)
    print(BACKUP)


if __name__ == "__main__":
    main()
